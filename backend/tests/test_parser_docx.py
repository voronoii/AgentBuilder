from __future__ import annotations

from pathlib import Path

import pytest

from app.services.knowledge.parsers.docx import DocxParser


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    from docx import Document as DocxDocument

    p = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("First paragraph with some body text.")
    doc.add_paragraph("Second paragraph.")
    doc.save(p)
    return p


@pytest.mark.asyncio
async def test_docx_parser_extracts_paragraphs(sample_docx: Path) -> None:
    parsed = await DocxParser().parse(sample_docx)
    assert "First paragraph" in parsed.text
    assert "Second paragraph" in parsed.text
