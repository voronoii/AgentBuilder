from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document as DocxDocument

from app.services.knowledge.parsers.base import ParsedDocument


class DocxParser:
    async def parse(self, path: Path) -> ParsedDocument:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, self._parse_sync, path)

    def _parse_sync(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts)
        return ParsedDocument(
            text=text,
            metadata={"paragraph_count": len(parts), "filename": path.name},
        )
